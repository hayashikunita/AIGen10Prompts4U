import streamlit as st
import json
from pathlib import Path
import random
from datetime import datetime
import os
from openai import OpenAI
from dotenv import load_dotenv
import pandas as pd
import io
import pdfplumber
from docx import Document

# 環境変数を読み込む
load_dotenv()

# ページ設定
st.set_page_config(
    page_title="AIGen10Prompts4U - システムプロンプト生成",
    page_icon="🤖",
    layout="wide"
)

# チャット履歴保存ディレクトリ
CHAT_HISTORY_DIR = Path("chat_history")
CHAT_HISTORY_DIR.mkdir(exist_ok=True)

# チャット履歴を保存する関数
def save_chat_history(title, messages, selected_prompt=None):
    """チャット履歴をJSONファイルに保存"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{title}_{timestamp}.json"
    filepath = CHAT_HISTORY_DIR / filename
    
    history_data = {
        "title": title,
        "timestamp": timestamp,
        "messages": messages,
        "selected_prompt": selected_prompt
    }
    
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(history_data, f, ensure_ascii=False, indent=2)
    
    return filepath

# チャット履歴一覧を取得する関数
def list_chat_histories():
    """保存されたチャット履歴のリストを取得"""
    histories = []
    for filepath in sorted(CHAT_HISTORY_DIR.glob("*.json"), reverse=True):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                histories.append({
                    "filename": filepath.name,
                    "filepath": filepath,
                    "title": data.get("title", "無題"),
                    "timestamp": data.get("timestamp", ""),
                    "message_count": len(data.get("messages", []))
                })
        except Exception as e:
            continue
    return histories

# チャット履歴を読み込む関数
def load_chat_history(filepath):
    """指定されたファイルからチャット履歴を読み込む"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        st.error(f"履歴の読み込みに失敗しました: {str(e)}")
        return None

# チャット履歴を削除する関数
def delete_chat_history(filepath):
    """指定されたチャット履歴を削除"""
    try:
        filepath.unlink()
        return True
    except Exception as e:
        st.error(f"履歴の削除に失敗しました: {str(e)}")
        return False

# OpenAI APIクライアント初期化
def get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY")
    if api_key:
        return OpenAI(api_key=api_key)
    return None

# トークン数を概算する関数（1トークン ≒ 4文字）
def estimate_tokens(text):
    """テキストのトークン数を概算"""
    return len(text) // 4

# ファイル内容を切り詰める関数
def truncate_content(content, max_tokens=15000):
    """
    コンテンツが大きすぎる場合に切り詰める
    max_tokens: 最大トークン数（デフォルト15000 ≒ 60,000文字）
    """
    estimated_tokens = estimate_tokens(content)
    
    if estimated_tokens <= max_tokens:
        return content, False  # 切り詰めなし
    
    # 切り詰める
    max_chars = max_tokens * 4
    truncated = content[:max_chars]
    
    # 最後の改行で切る（途中で切れないように）
    last_newline = truncated.rfind('\n')
    if last_newline > max_chars * 0.9:  # 90%以上の位置に改行があれば
        truncated = truncated[:last_newline]
    
    return truncated, True  # 切り詰めあり

# ファイル内容を読み取る関数
def read_file_content(uploaded_file):
    """
    アップロードされたファイルの内容を読み取り、テキスト形式で返す
    Excel、CSV、PDF、Word、テキストファイルに対応
    """
    file_name = uploaded_file.name
    file_extension = Path(file_name).suffix.lower()
    
    try:
        # PDFファイルの場合
        if file_extension == '.pdf':
            uploaded_file.seek(0)
            text_parts = []
            with pdfplumber.open(uploaded_file) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"\n--- ページ {i} ---\n{page_text}")
            
            if text_parts:
                content = "".join(text_parts)
                content += f"\n\n(総ページ数: {len(pdf.pages)})"
                return content, "pdf"
            else:
                return None, "error: PDFからテキストを抽出できませんでした"
        
        # Wordファイルの場合
        elif file_extension in ['.docx', '.doc']:
            uploaded_file.seek(0)
            if file_extension == '.docx':
                doc = Document(uploaded_file)
                text_parts = []
                for i, para in enumerate(doc.paragraphs, 1):
                    if para.text.strip():
                        text_parts.append(para.text)
                
                content = "\n".join(text_parts)
                if content.strip():
                    content += f"\n\n(段落数: {len([p for p in doc.paragraphs if p.text.strip()])})"
                    return content, "word"
                else:
                    return None, "error: Wordファイルからテキストを抽出できませんでした"
            else:
                return None, "error: .doc形式は非対応です。.docx形式に変換してください"
        
        # Excelファイルの場合
        elif file_extension in ['.xlsx', '.xls']:
            uploaded_file.seek(0)
            df_dict = pd.read_excel(uploaded_file, sheet_name=None)
            
            content_parts = []
            for sheet_name, df in df_dict.items():
                content_parts.append(f"\n=== シート: {sheet_name} ===\n")
                content_parts.append(df.to_string(index=False))
                content_parts.append(f"\n(行数: {len(df)}, 列数: {len(df.columns)})\n")
            
            return "".join(content_parts), "excel"
        
        # CSVファイルの場合
        elif file_extension == '.csv':
            uploaded_file.seek(0)
            for encoding in ['utf-8', 'shift_jis', 'cp932']:
                try:
                    uploaded_file.seek(0)
                    df = pd.read_csv(uploaded_file, encoding=encoding)
                    content = f"\n{df.to_string(index=False)}\n(行数: {len(df)}, 列数: {len(df.columns)})\n"
                    return content, "csv"
                except UnicodeDecodeError:
                    continue
            return None, "error: CSVファイルのエンコーディングを判別できませんでした"
        
        # テキストファイルの場合
        else:
            uploaded_file.seek(0)
            # 複数のエンコーディングを試す
            for encoding in ['utf-8', 'shift_jis', 'cp932', 'latin-1']:
                try:
                    uploaded_file.seek(0)
                    content = uploaded_file.read().decode(encoding)
                    return content, "text"
                except UnicodeDecodeError:
                    continue
            return None, "error: テキストファイルのエンコーディングを判別できませんでした"
            
    except Exception as e:
        return None, f"error: {str(e)}"

class PromptGenerator:
    def __init__(self):
        self.prompts_dir = Path("prompts_data")
        self.file_map = {
            "industry": "industry.json",
            "idea": "idea.json",
            "management": "management.json",
            "sales": "sales.json",
            "summary": "summary.json",
            "engineer": "engineer.json",
            "email": "email.json",
            "negotiation": "negotiation.json",
            "meeting": "meeting.json",
            "consultant": "consultant.json",
            "medical": "medical.json",
            "investment": "investment.json",
            "dating": "dating.json",
            "job_interview": "job_interview.json",
            "education": "education.json",
            "legal": "legal.json",
            "sns_content": "sns_content.json",
            "startup": "startup.json",
            "programmer": "programmer.json",
            "python_engineer": "python_engineer.json",
            "ai_engineer": "ai_engineer.json",
            "chatgpt_api": "chatgpt_api.json",
            "lawyer": "lawyer.json",
            "it_lawyer": "it_lawyer.json",
            "ceo": "ceo.json",
            "stock_trader": "stock_trader.json",
            "finance": "finance.json",
            "qol": "qol.json"
        }
        
        self.category_names = {
            "industry": "業界分析・市場調査用",
            "idea": "アイデア創出用",
            "management": "マネジメント用",
            "sales": "営業・セールス用",
            "summary": "要約・まとめ用",
            "engineer": "エンジニア用",
            "email": "メール返信用",
            "negotiation": "価格交渉用",
            "meeting": "会議準備用",
            "consultant": "コンサルティング用",
            "medical": "医療・健康相談用",
            "investment": "投資・資産運用用",
            "dating": "恋愛・デート用",
            "job_interview": "面接・転職対策用",
            "education": "教育・学習支援用",
            "legal": "法律・契約書用",
            "sns_content": "SNS・コンテンツ作成用",
            "startup": "起業・スタートアップ用",
            "programmer": "プログラマー実践用",
            "python_engineer": "Pythonエンジニア専門用",
            "ai_engineer": "AIエンジニア専門用",
            "chatgpt_api": "ChatGPT API活用専門用",
            "lawyer": "法律家・弁護士実践用",
            "it_lawyer": "IT法務・テック法律家専門用",
            "ceo": "経営者・CEO実践用",
            "stock_trader": "日本株トレーダー実践用",
            "finance": "金融業界・銀行実践用",
            "qol": "QOL向上・ライフスタイル改善用"
        }

    def load_prompts(self, category):
        """指定カテゴリのプロンプトを読み込む"""
        file_name = self.file_map.get(category)
        if not file_name:
            return None
        
        file_path = self.prompts_dir / file_name
        if not file_path.exists():
            return None
        
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return data

    def generate_prompts(self, category, count=10):
        """ランダムにプロンプトを生成"""
        data = self.load_prompts(category)
        if not data:
            return None
        
        prompts = data.get("prompts", [])
        if not prompts:
            return None
        
        # ランダムに選択
        selected = random.sample(prompts, min(count, len(prompts)))
        return {
            "category": data.get("category", category),
            "prompts": selected
        }

# Streamlitアプリ
def main():
    st.title("🤖 AIGen10Prompts4U")
    st.markdown("### システムプロンプト生成アプリ")
    st.markdown("---")
    
    generator = PromptGenerator()
    
    # セッション状態の初期化
    if "mode" not in st.session_state:
        st.session_state.mode = "generator"  # generator または chatbot
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "selected_prompt" not in st.session_state:
        st.session_state.selected_prompt = None
    
    # サイドバー
    with st.sidebar:
        st.header("⚙️ 設定")
        
        # モード選択
        current_mode = st.session_state.get("mode", "generator")
        mode = st.radio(
            "モードを選択",
            options=["generator", "chatbot"],
            format_func=lambda x: "🎲 プロンプト生成" if x == "generator" else "💬 チャットボット",
            index=0 if current_mode == "generator" else 1,
            key="mode_selector"
        )
        
        # モードが変更された場合のみ更新
        if mode != st.session_state.mode:
            st.session_state.mode = mode
        
        st.markdown("---")
        
        if st.session_state.mode == "generator":
            # カテゴリ選択
            category = st.selectbox(
                "カテゴリを選択",
                options=list(generator.file_map.keys()),
                format_func=lambda x: f"{x} ({generator.category_names.get(x, x)})"
            )
            
            # 生成数は最大値に固定(各カテゴリの全プロンプト)
            data = generator.load_prompts(category)
            count = len(data.get("prompts", [])) if data else 10
            st.info(f"📊 生成数: {count}個（全プロンプト）")
            
            # 生成ボタン
            generate_button = st.button("🎲 プロンプト生成", type="primary", use_container_width=True)
        else:
            # チャットボット設定
            st.markdown("**チャットボット設定**")
            
            # API キー確認
            client = get_openai_client()
            if client:
                st.success("✅ OpenAI API 接続済み")
            else:
                st.warning("⚠️ OpenAI APIキーが未設定です")
                st.markdown("`.env`ファイルに`OPENAI_API_KEY`を設定してください")
        
        st.markdown("---")
        st.markdown("**統計情報**")
        st.info(f"📊 合計22カテゴリ\n\n📝 合計780個のプロンプト")
    
    # メインエリア
    if st.session_state.mode == "generator":
        show_generator_mode(generator, generate_button if 'generate_button' in locals() else False, 
                          category if 'category' in locals() else None, 
                          count if 'count' in locals() else 10)
    else:
        show_chatbot_mode(generator)

def switch_to_chat(prompt):
    """チャットモードに切り替える"""
    st.session_state.selected_prompt = prompt
    st.session_state.messages = []
    st.session_state.mode = "chatbot"

def show_generator_mode(generator, generate_button, category, count):
    """プロンプト生成モード"""
    if generate_button:
        with st.spinner("プロンプトを生成中..."):
            result = generator.generate_prompts(category, count)
            
            if result:
                st.success(f"✅ {len(result['prompts'])}個のプロンプトを生成しました")
                st.markdown(f"**カテゴリ:** {result['category']}")
                st.markdown("---")
                
                # プロンプト一覧を表示
                st.markdown("### 📝 生成されたプロンプト")
                for i, prompt in enumerate(result['prompts'], 1):
                    with st.expander(f"**{i}. {prompt['title']}**"):
                        # システムプロンプト
                        st.markdown("##### 📋 システムプロンプト")
                        st.text_area(
                            "システムプロンプト",
                            value=prompt['system_prompt'],
                            height=150,
                            key=f"prompt_text_{i}_{hash(prompt['title'])}",
                            label_visibility="collapsed"
                        )
                        
                        # 推奨添付ファイル
                        st.markdown("##### 📎 推奨添付ファイル")
                        attachments = prompt.get('recommended_attachments', [])
                        if attachments:
                            for attachment in attachments:
                                st.markdown(f"• {attachment}")
                        else:
                            st.info("なし")
                        
                        # チャットボタン
                        st.markdown("---")
                        button_key = f"chat_expand_{i}_{hash(prompt['title'])}"
                        if st.button("💬 このプロンプトでチャット", key=button_key, type="primary", use_container_width=True, on_click=switch_to_chat, args=(prompt,)):
                            pass  # コールバックで処理
            else:
                st.error("❌ プロンプトの生成に失敗しました")
    else:
        # 初期表示
        st.info("👈 左のサイドバーでカテゴリと生成数を選択し、「プロンプト生成」ボタンをクリックしてください")
        
        # カテゴリ一覧を表示
        st.markdown("### 📚 利用可能なカテゴリ")
        
        col1, col2, col3 = st.columns(3)
        
        categories = list(generator.category_names.items())
        
        with col1:
            st.markdown("#### ビジネス系")
            for cat, name in categories[:8]:
                st.markdown(f"• **{cat}**: {name}")
        
        with col2:
            st.markdown("#### 実務・専門系")
            for cat, name in categories[8:15]:
                st.markdown(f"• **{cat}**: {name}")
        
        with col3:
            st.markdown("#### 技術系")
            for cat, name in categories[15:]:
                st.markdown(f"• **{cat}**: {name}")

def show_chatbot_mode(generator):
    """チャットボットモード"""
    client = get_openai_client()
    
    # 上部にボタンを配置
    col1, col2, col3, col4 = st.columns([2, 2, 2, 1])
    
    with col1:
        if st.button("📋 プロンプトから選択", use_container_width=True):
            st.session_state.show_prompt_selector = True
            st.rerun()
    
    with col2:
        if st.button("🆕 新しい会話", use_container_width=True):
            st.session_state.messages = []
            st.session_state.selected_prompt = None
            st.rerun()
    
    with col3:
        # 現在の会話を保存
        if st.button("💾 会話を保存", use_container_width=True, disabled=len(st.session_state.messages) == 0):
            st.session_state.show_save_dialog = True
    
    with col4:
        # 履歴を表示
        if st.button("📚", use_container_width=True, help="履歴を表示"):
            st.session_state.show_history = not st.session_state.get('show_history', False)
            st.rerun()
    
    # 会話保存ダイアログ
    if st.session_state.get('show_save_dialog', False):
        with st.expander("💾 会話を保存", expanded=True):
            save_title = st.text_input("タイトル", value=f"会話_{datetime.now().strftime('%Y%m%d_%H%M')}")
            
            col_save1, col_save2 = st.columns(2)
            with col_save1:
                if st.button("💾 保存", type="primary", use_container_width=True):
                    filepath = save_chat_history(
                        save_title,
                        st.session_state.messages,
                        st.session_state.selected_prompt
                    )
                    st.success(f"✅ 保存しました: {filepath.name}")
                    st.session_state.show_save_dialog = False
                    st.rerun()
            
            with col_save2:
                if st.button("❌ キャンセル", use_container_width=True):
                    st.session_state.show_save_dialog = False
                    st.rerun()
    
    # 履歴表示
    if st.session_state.get('show_history', False):
        with st.expander("📚 会話履歴", expanded=True):
            histories = list_chat_histories()
            
            if histories:
                st.markdown(f"**保存された会話: {len(histories)}件**")
                
                for hist in histories:
                    col_h1, col_h2, col_h3 = st.columns([3, 1, 1])
                    
                    with col_h1:
                        # タイトルと情報
                        timestamp_str = datetime.strptime(hist['timestamp'], "%Y%m%d_%H%M%S").strftime("%Y/%m/%d %H:%M")
                        st.markdown(f"**{hist['title']}**")
                        st.caption(f"📅 {timestamp_str} | 💬 {hist['message_count']}件のメッセージ")
                    
                    with col_h2:
                        # 読み込みボタン
                        if st.button("📂 読込", key=f"load_{hist['filename']}", use_container_width=True):
                            data = load_chat_history(hist['filepath'])
                            if data:
                                st.session_state.messages = data.get('messages', [])
                                st.session_state.selected_prompt = data.get('selected_prompt')
                                st.session_state.show_history = False
                                st.success(f"✅ {hist['title']} を読み込みました")
                                st.rerun()
                    
                    with col_h3:
                        # 削除ボタン
                        if st.button("🗑️", key=f"delete_{hist['filename']}", use_container_width=True, help="削除"):
                            if delete_chat_history(hist['filepath']):
                                st.success(f"✅ {hist['title']} を削除しました")
                                st.rerun()
                    
                    st.markdown("---")
            else:
                st.info("保存された会話はありません")
            
            if st.button("❌ 閉じる", use_container_width=True):
                st.session_state.show_history = False
                st.rerun()
    
    # プロンプト選択ダイアログ
    if hasattr(st.session_state, 'show_prompt_selector') and st.session_state.show_prompt_selector:
        with st.expander("📋 プロンプトを選択", expanded=True):
            category = st.selectbox(
                "カテゴリを選択",
                options=list(generator.file_map.keys()),
                format_func=lambda x: f"{generator.category_names.get(x, x)}",
                key="prompt_selector_category"
            )
            
            data = generator.load_prompts(category)
            if data:
                prompts = data.get("prompts", [])
                prompt_titles = [p['title'] for p in prompts]
                
                selected_title = st.selectbox(
                    "プロンプトを選択",
                    options=prompt_titles,
                    key="prompt_selector_title"
                )
                
                if st.button("✅ このプロンプトを使用", type="primary"):
                    selected_prompt = next((p for p in prompts if p['title'] == selected_title), None)
                    if selected_prompt:
                        st.session_state.selected_prompt = selected_prompt
                        st.session_state.messages = []
                        st.session_state.show_prompt_selector = False
                        st.rerun()
            
            if st.button("❌ キャンセル"):
                st.session_state.show_prompt_selector = False
                st.rerun()
    
    # 選択されたプロンプトを表示
    if st.session_state.selected_prompt:
        with st.expander("📋 使用中のシステムプロンプト", expanded=False):
            st.markdown(f"**タイトル:** {st.session_state.selected_prompt['title']}")
            st.text_area(
                "システムプロンプト",
                value=st.session_state.selected_prompt['system_prompt'],
                height=150,
                disabled=True,
                key="current_system_prompt"
            )
            
            # 推奨添付ファイルを表示
            st.markdown("**📎 推奨添付ファイル:**")
            attachments = st.session_state.selected_prompt.get('recommended_attachments', [])
            if attachments:
                for attachment in attachments:
                    st.markdown(f"• {attachment}")
            else:
                st.info("なし")
    
    # チャット履歴を表示
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            # 添付ファイル情報を表示
            if "files" in message and message["files"]:
                st.markdown("**📎 添付ファイル:**")
                for file_info in message["files"]:
                    st.markdown(f"• {file_info['name']} ({file_info['size']} bytes)")
    
    # ファイルアップローダー
    uploaded_files = st.file_uploader(
        "📎 ファイルを添付（複数可）",
        accept_multiple_files=True,
        key="file_uploader",
        help="PDF、Word (.docx)、Excel (.xlsx, .xls)、CSV、テキストファイル、コードなどに対応"
    )
    
    # ユーザー入力
    if prompt := st.chat_input("メッセージを入力してください..."):
        if not client:
            st.error("❌ OpenAI APIキーが設定されていません。`.env`ファイルに`OPENAI_API_KEY`を設定してください。")
            return
        
        # 添付ファイルの内容を読み取る
        file_contents = []
        file_info_list = []
        total_truncated = False
        
        if uploaded_files:
            for uploaded_file in uploaded_files:
                # 新しいファイル読み取り関数を使用
                content, file_type = read_file_content(uploaded_file)
                
                if content:
                    # コンテンツを切り詰める（1ファイルあたり最大15000トークン ≒ 60KB）
                    truncated_content, was_truncated = truncate_content(content, max_tokens=15000)
                    
                    if was_truncated:
                        total_truncated = True
                    
                    # 成功した場合
                    truncate_notice = " ⚠️ (ファイルが大きいため一部省略されました)" if was_truncated else ""
                    
                    if file_type == "pdf":
                        file_contents.append(f"\n\n--- 📕 {uploaded_file.name} (PDFファイル){truncate_notice} ---\n{truncated_content}")
                    elif file_type == "word":
                        file_contents.append(f"\n\n--- 📘 {uploaded_file.name} (Wordファイル){truncate_notice} ---\n{truncated_content}")
                    elif file_type == "excel":
                        file_contents.append(f"\n\n--- 📊 {uploaded_file.name} (Excelファイル){truncate_notice} ---\n{truncated_content}")
                    elif file_type == "csv":
                        file_contents.append(f"\n\n--- 📄 {uploaded_file.name} (CSVファイル){truncate_notice} ---\n{truncated_content}")
                    else:  # text
                        file_contents.append(f"\n\n--- 📝 {uploaded_file.name} (テキストファイル){truncate_notice} ---\n{truncated_content}")
                    
                    file_info_list.append({
                        "name": uploaded_file.name,
                        "size": uploaded_file.size,
                        "type": file_type,
                        "truncated": was_truncated
                    })
                else:
                    # 読み込みに失敗した場合
                    error_msg = f"⚠️ {uploaded_file.name} の読み込みに失敗しました"
                    if "error:" in file_type:
                        error_msg += f" ({file_type})"
                    file_contents.append(f"\n\n--- {error_msg} ---")
                    file_info_list.append({
                        "name": uploaded_file.name,
                        "size": uploaded_file.size,
                        "type": "error"
                    })
        
        # 警告メッセージを表示
        if total_truncated:
            st.warning("⚠️ 一部のファイルが大きすぎるため、内容の一部が省略されました。より詳細な分析が必要な場合は、ファイルを分割してアップロードしてください。")
        
        # プロンプトにファイル内容を追加
        full_prompt = prompt
        if file_contents:
            full_prompt += "\n\n" + "".join(file_contents)
        
        # 全体のトークン数をチェック
        total_tokens = estimate_tokens(full_prompt)
        if total_tokens > 25000:  # 25,000トークン以上の場合は警告
            st.error(f"❌ 入力が大きすぎます（推定 {total_tokens:,} トークン）。ファイルを分割するか、テキストを減らしてください。")
            return
        elif total_tokens > 20000:  # 20,000トークン以上の場合は注意喚起
            st.warning(f"⚠️ 入力が大きいです（推定 {total_tokens:,} トークン）。処理に時間がかかる可能性があります。")
        
        # ユーザーメッセージを追加
        user_message = {"role": "user", "content": full_prompt}
        if file_info_list:
            user_message["files"] = file_info_list
        
        st.session_state.messages.append(user_message)
        with st.chat_message("user"):
            st.markdown(prompt)
            if file_info_list:
                st.markdown("**📎 添付ファイル:**")
                for file_info in file_info_list:
                    st.markdown(f"• {file_info['name']} ({file_info['size']} bytes)")
        
        # アシスタントの応答を生成
        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            full_response = ""
            
            # システムプロンプトを含めてAPI呼び出し
            messages = []
            if st.session_state.selected_prompt:
                messages.append({
                    "role": "system",
                    "content": st.session_state.selected_prompt['system_prompt']
                })
            
            messages.extend([
                {"role": m["role"], "content": m["content"]}
                for m in st.session_state.messages
            ])
            
            try:
                stream = client.chat.completions.create(
                    # model="chatgpt-4o-latest",
                    model="gpt-5",
                    messages=messages,
                    stream=True,
                    # temperature=0.7
                )
                
                for chunk in stream:
                    if chunk.choices[0].delta.content is not None:
                        full_response += chunk.choices[0].delta.content
                        message_placeholder.markdown(full_response + "▌")
                
                message_placeholder.markdown(full_response)
                
            except Exception as e:
                full_response = f"❌ エラーが発生しました: {str(e)}"
                message_placeholder.markdown(full_response)
            
            # アシスタントメッセージを追加
            st.session_state.messages.append({"role": "assistant", "content": full_response})
    
    # 初期メッセージ
    if len(st.session_state.messages) == 0:
        st.info("💬 チャットを開始してください。左のサイドバーから「プロンプトから選択」でシステムプロンプトを設定できます。")

if __name__ == "__main__":
    main()
